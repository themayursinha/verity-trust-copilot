from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.dependencies import get_current_active_user
from app.models.answer import Answer, AnswerGeneration
from app.models.user import User

ROOT = Path(__file__).resolve().parent.parent.parent.parent
OUTPUT_DIR = ROOT / "outputs"

router = APIRouter(prefix="/api/v1/export", tags=["export"])


class ExportAnswerRequest(BaseModel):
    answer: dict[str, Any]


class ExportQuestionnaireRequest(BaseModel):
    generation_id: str
    format: str = "xlsx"


def render_customer_ready_markdown(answer: dict[str, Any]) -> str:
    lines = [
        "# Customer Security Answer",
        "",
        "## Question",
        "",
        str(answer.get("question", "")),
        "",
        "## Draft Answer",
        "",
        str(answer.get("answer", "")),
        "",
        f"**Confidence:** {answer.get('confidence', 'unknown')}",
        f"**Needs human review:** {str(answer.get('needs_human_review', False)).lower()}",
        "",
        "## Sources",
        "",
    ]
    citations = answer.get("citations") if isinstance(answer.get("citations"), list) else []
    if citations:
        for citation in citations:
            if not isinstance(citation, dict):
                continue
            lines.append(
                "- {citation} {title} ({source_id}, reviewed {last_reviewed})".format(
                    citation=citation.get("citation", ""),
                    title=citation.get("title", ""),
                    source_id=citation.get("source_id", ""),
                    last_reviewed=citation.get("last_reviewed", ""),
                )
            )
    else:
        lines.append("- No sources found. Do not send without review.")
    lines.extend(["", "## Freshness", ""])
    freshness = answer.get("freshness") if isinstance(answer.get("freshness"), list) else []
    if freshness:
        for item in freshness:
            if not isinstance(item, dict):
                continue
            lines.append(
                "- {source}: {status}, reviewed {last_reviewed}, {age_days} days old".format(
                    source=item.get("source", ""),
                    status=item.get("status", ""),
                    last_reviewed=item.get("last_reviewed", ""),
                    age_days=item.get("age_days", ""),
                )
            )
    else:
        lines.append("- No freshness checks available.")
    lines.append("")
    return "\n".join(lines)


def generate_export_xlsx(generation_id: str, answers: list[Answer]) -> BytesIO:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Security Responses"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="2B579A", end_color="2B579A", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    headers = ["#", "Question", "Answer", "Confidence", "Needs Review", "Source"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    for row_idx, answer in enumerate(answers, 2):
        values = [
            row_idx - 1,
            answer.question,
            answer.answer_text,
            answer.confidence or "unknown",
            "Yes" if answer.needs_human_review else "No",
            answer.source or "ai",
        ]
        for col, value in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = thin_border

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 60
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 14
    ws.column_dimensions["F"].width = 10

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def generate_export_docx(answers: list[Answer]) -> BytesIO:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("Security Questionnaire Responses", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph(f"Total questions: {len(answers)}")
    doc.add_paragraph("")

    for i, answer in enumerate(answers, 1):
        doc.add_heading(f"Question {i}", level=2)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(answer.question)
        q_run.bold = True
        q_run.font.size = Pt(11)

        doc.add_heading("Answer", level=3)
        doc.add_paragraph(answer.answer_text)

        doc.add_heading("Metadata", level=3)
        doc.add_paragraph(f"Confidence: {answer.confidence or 'unknown'}")
        doc.add_paragraph(f"Needs human review: {'Yes' if answer.needs_human_review else 'No'}")
        doc.add_paragraph(f"Source: {answer.source or 'ai'}")

        if answer.citations:
            doc.add_heading("Sources", level=4)
            for citation in answer.citations:
                if isinstance(citation, dict):
                    doc.add_paragraph(
                        f"{citation.get('citation', '')} {citation.get('title', '')} "
                        f"(ID: {citation.get('source_id', '')})",
                        style="List Bullet",
                    )

        doc.add_paragraph("_" * 40)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


@router.post("/", response_model=dict)
async def export_markdown(
    body: ExportAnswerRequest,
    current_user: User = Depends(get_current_active_user),
):
    answer = body.answer
    if not isinstance(answer, dict):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Send an answer object to export.")
    markdown = render_customer_ready_markdown(answer)
    return {"path": str(OUTPUT_DIR / "customer_ready_answer.md"), "markdown": markdown}


@router.post("/csv", response_model=dict)
async def export_csv_answer(
    body: ExportAnswerRequest,
    current_user: User = Depends(get_current_active_user),
):
    from app.core.engine import export_csv

    answer = body.answer
    if not isinstance(answer, dict):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Send an answer object to export.")
    results = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "as_of_date": date.today().isoformat(),
        "summary": {
            "questions_processed": 1,
            "confidence_counts": {answer.get("confidence", "low"): 1},
            "human_reviews_required": 1 if answer.get("needs_human_review") else 0,
        },
        "answers": [answer],
    }
    csv_content = export_csv(results)
    return {"path": str(OUTPUT_DIR / "customer_ready_answer.csv"), "csv": csv_content}


@router.post("/json", response_model=dict)
async def export_json_answer(
    body: ExportAnswerRequest,
    current_user: User = Depends(get_current_active_user),
):
    answer = body.answer
    if not isinstance(answer, dict):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Send an answer object to export.")
    results = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "as_of_date": date.today().isoformat(),
        "summary": {
            "questions_processed": 1,
            "confidence_counts": {answer.get("confidence", "low"): 1},
            "human_reviews_required": 1 if answer.get("needs_human_review") else 0,
        },
        "answers": [answer],
    }
    return {"path": str(OUTPUT_DIR / "customer_ready_answer.json"), "json": results}


@router.post("/questionnaire/xlsx")
async def export_questionnaire_xlsx(
    body: ExportQuestionnaireRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    gen_result = await db.execute(
        select(AnswerGeneration).where(
            AnswerGeneration.id == body.generation_id,
            AnswerGeneration.org_id == current_user.org_id,
        )
    )
    gen = gen_result.scalar_one_or_none()
    if not gen:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Answer generation not found")

    ans_result = await db.execute(select(Answer).where(Answer.generation_id == gen.id).order_by(Answer.order_index))
    answers = ans_result.scalars().all()

    xlsx_buffer = generate_export_xlsx(gen.id, answers)

    filename = gen.original_filename or f"questionnaire_{gen.id[:8]}.xlsx"
    if not filename.endswith(".xlsx"):
        filename = f"{filename.rsplit('.', 1)[0]}.xlsx"

    return StreamingResponse(
        xlsx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/questionnaire/docx")
async def export_questionnaire_docx(
    body: ExportQuestionnaireRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    gen_result = await db.execute(
        select(AnswerGeneration).where(
            AnswerGeneration.id == body.generation_id,
            AnswerGeneration.org_id == current_user.org_id,
        )
    )
    gen = gen_result.scalar_one_or_none()
    if not gen:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Answer generation not found")

    ans_result = await db.execute(select(Answer).where(Answer.generation_id == gen.id).order_by(Answer.order_index))
    answers = ans_result.scalars().all()

    docx_buffer = generate_export_docx(answers)

    filename = gen.original_filename or f"questionnaire_{gen.id[:8]}.docx"
    if not filename.endswith(".docx"):
        filename = f"{filename.rsplit('.', 1)[0]}.docx"

    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
