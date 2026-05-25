import re
from io import BytesIO

from fastapi import UploadFile


ALLOWED_EXTENSIONS = {".xlsx", ".docx"}
MAX_FILE_SIZE = 10 * 1024 * 1024


def _parse_questions_from_text(text: str) -> list[str]:
    if not text.strip():
        return []

    lines = text.split("\n")
    questions: list[str] = []
    current = ""

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if current:
                questions.append(current.strip())
                current = ""
            continue

        is_new = (
            bool(re.match(r"^\d+[.)]\s", line))
            or bool(re.match(r"^Q\d+[:.]\s", line, re.IGNORECASE))
            or bool(re.match(r"^Question\s*\d+[:.]\s", line, re.IGNORECASE))
            or bool(re.match(r"^[•\-*]\s", line))
            or bool(re.match(r"^Q:\s", line, re.IGNORECASE))
            or bool(re.match(r"^[A-Z][.)]\s", line))
            or line.rstrip().endswith("?")
            or bool(re.match(r"^\d+\.\d+\s", line))
        )

        if is_new:
            if current:
                questions.append(current.strip())
            current = re.sub(r"^\d+[.)]\s*", "", line)
            current = re.sub(r"^Q\d+[:.]\s*", "", current, flags=re.IGNORECASE)
            current = re.sub(r"^Question\s*\d+[:.]\s*", "", current, flags=re.IGNORECASE)
            current = re.sub(r"^[•\-*]\s*", "", current)
            current = re.sub(r"^Q:\s*", "", current, flags=re.IGNORECASE)
            current = re.sub(r"^[A-Z][.)]\s*", "", current)
        elif current:
            current += " " + line
        else:
            if len(line) > 10 and (
                "?" in line
                or bool(
                    re.match(
                        r"^(Do|Does|Is|Are|Can|Will|Has|Have|What|How|When|Where|Who|Why|Which|Please|Describe|Explain|List|Provide|Outline)",
                        line,
                        re.IGNORECASE,
                    )
                )
            ):
                current = line

    if current:
        questions.append(current.strip())

    return [q for q in questions if len(q) > 5]


async def parse_questionnaire_file(file: UploadFile) -> list[str]:
    if not file.filename:
        raise ValueError("No filename provided")

    ext = "." + file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Please upload .xlsx or .docx files.")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise ValueError(f"File too large. Maximum size is {MAX_FILE_SIZE // (1024 * 1024)} MB.")

    if ext == ".xlsx":
        text = _parse_xlsx(content)
    else:
        text = _parse_docx(content)

    return _parse_questions_from_text(text)


def _parse_xlsx(content: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(BytesIO(content), read_only=True, data_only=True)
    lines: list[str] = []
    header_keywords = {
        "category",
        "question",
        "domain",
        "section",
        "topic",
        "control",
        "requirement",
        "status",
        "response",
        "#",
        "no.",
        "number",
        "id",
        "ref",
        "reference",
    }

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        first_row = True
        for row in ws.iter_rows():
            row_values = [str(cell.value).strip() for cell in row if cell.value is not None and str(cell.value).strip()]

            if not row_values:
                continue

            joined = " ".join(row_values)

            if first_row:
                first_row = False
                lower_vals = {v.lower() for v in row_values}
                if lower_vals & header_keywords:
                    continue

            lines.append(joined)

    wb.close()
    return "\n".join(lines)


def _parse_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(content))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)
